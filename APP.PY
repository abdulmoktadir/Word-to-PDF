import io
import json
import math
import platform
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Optional, Tuple

import fitz  # PyMuPDF
import streamlit as st
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt
from pdf2docx import Converter

# ==================================================
# PAGE CONFIG
# ==================================================
st.set_page_config(
    page_title="PDF ↔ Word Converter Pro",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ==================================================
# STYLES
# ==================================================
st.markdown(
    """
    <style>
        .stApp {
            background: linear-gradient(180deg, #f8fbff 0%, #eef4ff 100%);
        }
        .block-container {
            max-width: 1180px;
            padding-top: 1.4rem;
            padding-bottom: 2rem;
        }
        .hero {
            background: linear-gradient(135deg, #0f172a 0%, #1d4ed8 58%, #2563eb 100%);
            padding: 2rem 2rem 1.8rem 2rem;
            border-radius: 28px;
            color: white;
            box-shadow: 0 22px 45px rgba(37, 99, 235, 0.22);
            border: 1px solid rgba(255,255,255,0.14);
            margin-bottom: 1rem;
        }
        .hero h1 {
            margin: 0 0 0.3rem 0;
            font-size: 2.35rem;
            line-height: 1.12;
        }
        .hero p {
            color: rgba(255,255,255,0.93);
            font-size: 1rem;
            line-height: 1.8;
            max-width: 900px;
            margin-bottom: 0;
        }
        .tag {
            display: inline-block;
            background: rgba(255,255,255,0.15);
            color: white;
            border: 1px solid rgba(255,255,255,0.18);
            padding: 0.35rem 0.8rem;
            border-radius: 999px;
            margin-right: 0.45rem;
            margin-top: 0.45rem;
            font-size: 0.86rem;
            font-weight: 600;
        }
        .soft-card {
            background: #ffffff;
            border: 1px solid #e2e8f0;
            border-radius: 22px;
            padding: 1.05rem 1.05rem;
            box-shadow: 0 10px 24px rgba(15, 23, 42, 0.06);
            margin-bottom: 1rem;
        }
        .glass-card {
            background: rgba(255,255,255,0.85);
            border: 1px solid rgba(226,232,240,0.85);
            border-radius: 24px;
            padding: 1.15rem 1.15rem;
            box-shadow: 0 12px 28px rgba(15, 23, 42, 0.07);
            backdrop-filter: blur(6px);
            margin-bottom: 1rem;
        }
        .mini-card {
            background: white;
            border: 1px solid #dbeafe;
            border-radius: 18px;
            padding: 0.95rem;
            box-shadow: 0 8px 18px rgba(37, 99, 235, 0.06);
            height: 100%;
        }
        .section-title {
            font-size: 1.3rem;
            font-weight: 700;
            color: #0f172a;
            margin-bottom: 0.35rem;
        }
        .muted {
            color: #64748b;
            font-size: 0.96rem;
            line-height: 1.75;
        }
        .blue-tag {
            display: inline-block;
            background: #dbeafe;
            color: #1d4ed8;
            padding: 0.28rem 0.72rem;
            border-radius: 999px;
            margin-right: 0.45rem;
            margin-bottom: 0.45rem;
            font-size: 0.84rem;
            font-weight: 600;
        }
        .status-good {
            background: #ecfdf5;
            border: 1px solid #a7f3d0;
            color: #065f46;
            border-radius: 16px;
            padding: 0.85rem 1rem;
            font-weight: 600;
            margin-bottom: 1rem;
        }
        .status-warn {
            background: #fff7ed;
            border: 1px solid #fdba74;
            color: #9a3412;
            border-radius: 16px;
            padding: 0.85rem 1rem;
            font-weight: 600;
            margin-bottom: 1rem;
        }
        .status-neutral {
            background: #eff6ff;
            border: 1px solid #bfdbfe;
            color: #1d4ed8;
            border-radius: 16px;
            padding: 0.85rem 1rem;
            font-weight: 600;
            margin-bottom: 1rem;
        }
        .file-box {
            background: linear-gradient(180deg, #ffffff 0%, #f8fbff 100%);
            border: 1px solid #dbeafe;
            border-radius: 18px;
            padding: 1rem;
            margin-bottom: 0.9rem;
        }
        .footer-note {
            text-align: center;
            color: #64748b;
            padding-top: 1rem;
            font-size: 0.95rem;
        }
        .stButton > button,
        .stDownloadButton > button {
            border-radius: 14px !important;
            border: none !important;
            padding: 0.68rem 1rem !important;
            font-weight: 700 !important;
            box-shadow: 0 10px 18px rgba(37, 99, 235, 0.15) !important;
        }
        .stDownloadButton > button {
            width: 100%;
        }
        [data-testid="stFileUploader"] {
            background: rgba(255,255,255,0.72);
            border: 1px dashed #93c5fd;
            border-radius: 18px;
            padding: 0.35rem;
        }
        [data-testid="stSidebar"] {
            background: linear-gradient(180deg, #f8fbff 0%, #eef4ff 100%);
        }
    </style>
    """,
    unsafe_allow_html=True,
)

# ==================================================
# HELPERS
# ==================================================
def write_uploaded_file(uploaded_file, destination: Path) -> Path:
    destination.write_bytes(uploaded_file.getbuffer())
    return destination


def readable_size(size_bytes: int) -> str:
    units = ["B", "KB", "MB", "GB"]
    size = float(size_bytes)
    for unit in units:
        if size < 1024 or unit == units[-1]:
            return f"{size:.2f} {unit}"
        size /= 1024
    return f"{size_bytes} B"


def show_file_info(uploaded_file, expected_output: str) -> None:
    st.markdown(
        f"""
        <div class="file-box">
            <div style="font-weight:700; color:#0f172a; margin-bottom:0.35rem;">Selected file</div>
            <div class="muted"><b>Name:</b> {uploaded_file.name}</div>
            <div class="muted"><b>Size:</b> {readable_size(uploaded_file.size)}</div>
            <div class="muted"><b>Output:</b> {expected_output}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def convert_pdf_to_docx_editable(input_pdf: Path, output_docx: Path) -> None:
    converter = Converter(str(input_pdf))
    try:
        converter.convert(str(output_docx))
    finally:
        converter.close()


def set_section_size_from_pdf(section, page_rect) -> None:
    # PDF points and Word points are both 1/72 inch.
    section.page_width = Pt(page_rect.width)
    section.page_height = Pt(page_rect.height)
    # Small margins so the image nearly fills the page.
    section.top_margin = Inches(0.25)
    section.bottom_margin = Inches(0.25)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)


def convert_pdf_to_docx_preserve_appearance(
    input_pdf: Path,
    output_docx: Path,
    zoom_factor: float = 2.0,
) -> None:
    """
    Preserve page appearance by rendering each PDF page as an image
    and inserting it into a DOCX page. This keeps visual fidelity higher,
    but the result is not truly editable text.
    """
    pdf = fitz.open(str(input_pdf))
    doc = Document()

    # Remove the default empty paragraph spacing a bit.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    first_page = True
    for page_index, page in enumerate(pdf):
        rect = page.rect
        if first_page:
            section = doc.sections[0]
            set_section_size_from_pdf(section, rect)
            first_page = False
        else:
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            set_section_size_from_pdf(section, rect)

        matrix = fitz.Matrix(zoom_factor, zoom_factor)
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        img_bytes = pix.tobytes("png")

        usable_width_inches = max((rect.width / 72.0) - 0.5, 1.0)
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run()
        run.add_picture(io.BytesIO(img_bytes), width=Inches(usable_width_inches))

    doc.save(str(output_docx))
    pdf.close()


def build_libreoffice_filter_json(
    lossless: bool = True,
    quality: int = 100,
    reduce_resolution: bool = False,
    max_resolution: int = 300,
) -> str:
    options = {
        "UseLosslessCompression": {"type": "boolean", "value": str(lossless).lower()},
        "Quality": {"type": "long", "value": str(quality)},
        "ReduceImageResolution": {"type": "boolean", "value": str(reduce_resolution).lower()},
        "MaxImageResolution": {"type": "long", "value": str(max_resolution)},
    }
    return json.dumps(options, separators=(",", ":"))


def try_libreoffice_docx_to_pdf(
    input_docx: Path,
    output_dir: Path,
    high_quality: bool = True,
) -> Tuple[bool, str, Optional[Path]]:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return False, "LibreOffice was not found on this machine.", None

    profile_dir = output_dir / "lo_profile"
    profile_dir.mkdir(parents=True, exist_ok=True)
    profile_uri = profile_dir.resolve().as_uri()

    if high_quality:
        filter_json = build_libreoffice_filter_json(
            lossless=True,
            quality=100,
            reduce_resolution=False,
            max_resolution=300,
        )
        convert_arg = f"pdf:writer_pdf_Export:{filter_json}"
    else:
        convert_arg = "pdf"

    cmd = [
        soffice,
        f"-env:UserInstallation={profile_uri}",
        "--headless",
        "--convert-to",
        convert_arg,
        "--outdir",
        str(output_dir),
        str(input_docx),
    ]

    result = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        timeout=240,
    )

    output_pdf = output_dir / f"{input_docx.stem}.pdf"
    if result.returncode == 0 and output_pdf.exists():
        message = (result.stdout or "DOCX converted to PDF successfully.").strip()
        return True, message, output_pdf

    error_message = (result.stderr or result.stdout or "LibreOffice conversion failed.").strip()
    return False, error_message, None


def try_docx2pdf_fallback(input_docx: Path, output_pdf: Path) -> Tuple[bool, str, Optional[Path]]:
    system_name = platform.system().lower()
    if system_name not in {"windows", "darwin"}:
        return False, "docx2pdf fallback is only useful on Windows or macOS.", None

    try:
        from docx2pdf import convert as docx2pdf_convert
    except Exception:
        return False, "docx2pdf is not installed.", None

    try:
        docx2pdf_convert(str(input_docx), str(output_pdf))
        if output_pdf.exists():
            return True, "Converted using Microsoft Word via docx2pdf.", output_pdf
        return False, "docx2pdf finished but no PDF was created.", None
    except Exception as exc:
        return False, f"docx2pdf failed: {exc}", None


def convert_docx_to_pdf(
    input_docx: Path,
    output_dir: Path,
    high_quality: bool = True,
) -> Tuple[bool, str, Optional[Path]]:
    ok, message, pdf_path = try_libreoffice_docx_to_pdf(
        input_docx,
        output_dir,
        high_quality=high_quality,
    )
    if ok:
        return ok, message, pdf_path

    fallback_pdf = output_dir / f"{input_docx.stem}.pdf"
    ok2, message2, pdf_path2 = try_docx2pdf_fallback(input_docx, fallback_pdf)
    if ok2:
        return ok2, message2, pdf_path2

    combined = (
        "DOCX → PDF could not be completed. "
        f"LibreOffice result: {message} | docx2pdf result: {message2}"
    )
    return False, combined, None


# ==================================================
# SIDEBAR
# ==================================================
st.sidebar.title("Converter Panel")
st.sidebar.markdown("### Supported modes")
st.sidebar.markdown("- PDF → DOCX (editable)")
st.sidebar.markdown("- PDF → DOCX (preserve appearance)")
st.sidebar.markdown("- DOCX → PDF (high quality)")

st.sidebar.markdown("---")
st.sidebar.markdown("### Privacy")
st.sidebar.markdown(
    "Files are processed temporarily for conversion. This app does not intentionally save them permanently in its own logic."
)

st.sidebar.markdown("---")
st.sidebar.markdown("### Deployment files")
st.sidebar.code("""app.py\nrequirements.txt\npackages.txt""")

st.sidebar.markdown("### requirements.txt")
st.sidebar.code(
    """streamlit
pdf2docx
python-docx
PyMuPDF
docx2pdf"""
)

st.sidebar.markdown("### packages.txt")
st.sidebar.code("""libreoffice""")

# ==================================================
# HERO
# ==================================================
st.markdown(
    """
    <div class="hero">
        <h1>PDF ↔ Word Converter Pro</h1>
        <p>
            A polished Streamlit converter with a cleaner interface, editable conversion mode,
            appearance-preserving conversion mode, and high-quality DOCX to PDF export.
            This version is better suited for documents that contain important images.
        </p>
        <div>
            <span class="tag">Modern UI</span>
            <span class="tag">Editable Mode</span>
            <span class="tag">Preserve Appearance Mode</span>
            <span class="tag">High-Quality Export</span>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)

c1, c2, c3 = st.columns(3)
with c1:
    st.markdown(
        """
        <div class="mini-card">
            <div style="font-size:1.4rem; font-weight:800; color:#0f172a;">01</div>
            <div style="font-weight:700; margin-top:0.2rem;">Choose mode</div>
            <div class="muted">Use editable conversion for normal documents, or preserve appearance when image fidelity matters more.</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
with c2:
    st.markdown(
        """
        <div class="mini-card">
            <div style="font-size:1.4rem; font-weight:800; color:#0f172a;">02</div>
            <div style="font-weight:700; margin-top:0.2rem;">Convert safely</div>
            <div class="muted">Files are handled in a temporary working folder during conversion and then cleaned up.</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
with c3:
    st.markdown(
        """
        <div class="mini-card">
            <div style="font-size:1.4rem; font-weight:800; color:#0f172a;">03</div>
            <div style="font-weight:700; margin-top:0.2rem;">Download</div>
            <div class="muted">Get your converted file immediately after a successful run.</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

soffice_available = shutil.which("soffice") or shutil.which("libreoffice")
if soffice_available:
    st.markdown(
        f"<div class='status-good'>✅ LibreOffice detected. DOCX → PDF high-quality export is available through: {soffice_available}</div>",
        unsafe_allow_html=True,
    )
else:
    st.markdown(
        "<div class='status-warn'>⚠️ LibreOffice is not detected right now. PDF → DOCX will still work, but DOCX → PDF on Streamlit Cloud needs <b>libreoffice</b> in <b>packages.txt</b>.</div>",
        unsafe_allow_html=True,
    )

# ==================================================
# TABS
# ==================================================
tab1, tab2, tab3 = st.tabs([
    "📄 PDF to Word (Editable)",
    "🖼️ PDF to Word (Preserve Appearance)",
    "📝 Word to PDF (High Quality)",
])

# ==================================================
# TAB 1: PDF -> WORD EDITABLE
# ==================================================
with tab1:
    left, right = st.columns([1.2, 0.8], gap="large")

    with left:
        st.markdown("<div class='glass-card'>", unsafe_allow_html=True)
        st.markdown("<div class='section-title'>Convert PDF into editable DOCX</div>", unsafe_allow_html=True)
        st.markdown(
            "<div class='muted'>This mode tries to preserve text, paragraphs, and layout in an editable Word file. It is best for normal text-based PDFs.</div>",
            unsafe_allow_html=True,
        )

        uploaded_pdf_edit = st.file_uploader(
            "Upload a PDF file",
            type=["pdf"],
            key="pdf_uploader_editable",
        )

        if uploaded_pdf_edit is not None:
            default_docx_name = Path(uploaded_pdf_edit.name).stem + "_editable.docx"
            show_file_info(uploaded_pdf_edit, default_docx_name)

            custom_docx_name = st.text_input(
                "Output DOCX file name",
                value=default_docx_name,
                key="pdf_to_docx_edit_name",
            )

            if st.button("Convert PDF to editable Word", use_container_width=True, key="btn_pdf_editable"):
                with tempfile.TemporaryDirectory() as tmpdir:
                    tmpdir_path = Path(tmpdir)
                    input_pdf = tmpdir_path / uploaded_pdf_edit.name
                    output_docx = tmpdir_path / custom_docx_name
                    write_uploaded_file(uploaded_pdf_edit, input_pdf)

                    try:
                        with st.spinner("Converting PDF to editable DOCX..."):
                            convert_pdf_to_docx_editable(input_pdf, output_docx)

                        if output_docx.exists():
                            st.success("Conversion completed successfully.")
                            st.download_button(
                                label="Download editable DOCX",
                                data=output_docx.read_bytes(),
                                file_name=custom_docx_name,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key="download_editable_docx",
                            )
                        else:
                            st.error("Conversion finished but the DOCX file was not created.")
                    except Exception as exc:
                        st.error(f"Conversion failed: {exc}")

        st.markdown("</div>", unsafe_allow_html=True)

    with right:
        st.markdown("<div class='soft-card'>", unsafe_allow_html=True)
        st.markdown("<div class='section-title'>Best for</div>", unsafe_allow_html=True)
        st.markdown("<span class='blue-tag'>Articles</span><span class='blue-tag'>Reports</span><span class='blue-tag'>Normal PDFs</span>", unsafe_allow_html=True)
        st.markdown(
            "<div class='muted'>Choose this when editability is more important than exact visual matching.</div>",
            unsafe_allow_html=True,
        )
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<div class='soft-card'>", unsafe_allow_html=True)
        st.markdown("<div class='section-title'>Notes</div>", unsafe_allow_html=True)
        st.markdown(
            """
            <div class='muted'>
            • Good for text-based PDFs.<br>
            • Very complex images, vector graphics, and equations may shift.<br>
            • Review the output before final use.
            </div>
            """,
            unsafe_allow_html=True,
        )
        st.markdown("</div>", unsafe_allow_html=True)

# ==================================================
# TAB 2: PDF -> WORD PRESERVE APPEARANCE
# ==================================================
with tab2:
    left2, right2 = st.columns([1.2, 0.8], gap="large")

    with left2:
        st.markdown("<div class='glass-card'>", unsafe_allow_html=True)
        st.markdown("<div class='section-title'>Convert PDF into Word while preserving page appearance</div>", unsafe_allow_html=True)
        st.markdown(
            "<div class='muted'>This mode renders each PDF page as a high-resolution image and inserts it into Word. It usually preserves the original visual appearance better, especially for image-heavy pages, but the result is not truly editable text.</div>",
            unsafe_allow_html=True,
        )

        uploaded_pdf_visual = st.file_uploader(
            "Upload a PDF file",
            type=["pdf"],
            key="pdf_uploader_visual",
        )

        zoom_factor = st.slider(
            "Render quality",
            min_value=1.5,
            max_value=4.0,
            value=2.5,
            step=0.5,
            help="Higher values preserve appearance better but increase file size and conversion time.",
        )

        if uploaded_pdf_visual is not None:
            default_docx_visual_name = Path(uploaded_pdf_visual.name).stem + "_appearance.docx"
            show_file_info(uploaded_pdf_visual, default_docx_visual_name)

            custom_visual_docx_name = st.text_input(
                "Output DOCX file name",
                value=default_docx_visual_name,
                key="pdf_to_docx_visual_name",
            )

            if st.button("Convert PDF to appearance-preserving Word", use_container_width=True, key="btn_pdf_visual"):
                with tempfile.TemporaryDirectory() as tmpdir:
                    tmpdir_path = Path(tmpdir)
                    input_pdf = tmpdir_path / uploaded_pdf_visual.name
                    output_docx = tmpdir_path / custom_visual_docx_name
                    write_uploaded_file(uploaded_pdf_visual, input_pdf)

                    try:
                        with st.spinner("Rendering pages and building DOCX..."):
                            convert_pdf_to_docx_preserve_appearance(
                                input_pdf,
                                output_docx,
                                zoom_factor=zoom_factor,
                            )

                        if output_docx.exists():
                            st.success("Appearance-preserving DOCX created successfully.")
                            st.download_button(
                                label="Download appearance-preserving DOCX",
                                data=output_docx.read_bytes(),
                                file_name=custom_visual_docx_name,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key="download_visual_docx",
                            )
                        else:
                            st.error("Conversion finished but the DOCX file was not created.")
                    except Exception as exc:
                        st.error(f"Conversion failed: {exc}")

        st.markdown("</div>", unsafe_allow_html=True)

    with right2:
        st.markdown("<div class='soft-card'>", unsafe_allow_html=True)
        st.markdown("<div class='section-title'>Best for</div>", unsafe_allow_html=True)
        st.markdown("<span class='blue-tag'>Image-heavy PDFs</span><span class='blue-tag'>Scanned layout</span><span class='blue-tag'>Exact appearance</span>", unsafe_allow_html=True)
        st.markdown(
            "<div class='muted'>Choose this when preserving the original look matters more than editability.</div>",
            unsafe_allow_html=True,
        )
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<div class='soft-card'>", unsafe_allow_html=True)
        st.markdown("<div class='section-title'>Trade-off</div>", unsafe_allow_html=True)
        st.markdown(
            """
            <div class='muted'>
            • Better visual fidelity for images and graphics.<br>
            • Output file can be larger.<br>
            • Text on the page behaves like an image, not editable paragraphs.
            </div>
            """,
            unsafe_allow_html=True,
        )
        st.markdown("</div>", unsafe_allow_html=True)

# ==================================================
# TAB 3: WORD -> PDF HIGH QUALITY
# ==================================================
with tab3:
    left3, right3 = st.columns([1.2, 0.8], gap="large")

    with left3:
        st.markdown("<div class='glass-card'>", unsafe_allow_html=True)
        st.markdown("<div class='section-title'>Convert DOCX into high-quality PDF</div>", unsafe_allow_html=True)
        st.markdown(
            "<div class='muted'>This mode uses LibreOffice in headless mode with quality-focused PDF export settings, including lossless compression and no image downsampling.</div>",
            unsafe_allow_html=True,
        )

        uploaded_docx = st.file_uploader(
            "Upload a DOCX file",
            type=["docx"],
            key="docx_uploader_high_quality",
        )

        high_quality = st.toggle(
            "Use high-quality image preservation",
            value=True,
            help="Keeps lossless compression and avoids image downsampling when LibreOffice is used.",
        )

        if uploaded_docx is not None:
            default_pdf_name = Path(uploaded_docx.name).stem + ".pdf"
            show_file_info(uploaded_docx, default_pdf_name)

            custom_pdf_name = st.text_input(
                "Output PDF file name",
                value=default_pdf_name,
                key="docx_to_pdf_name_hq",
            )

            if st.button("Convert Word to PDF", use_container_width=True, key="btn_docx_to_pdf_hq"):
                with tempfile.TemporaryDirectory() as tmpdir:
                    tmpdir_path = Path(tmpdir)
                    input_docx = tmpdir_path / uploaded_docx.name
                    write_uploaded_file(uploaded_docx, input_docx)

                    try:
                        with st.spinner("Converting DOCX to PDF..."):
                            ok, message, pdf_path = convert_docx_to_pdf(
                                input_docx,
                                tmpdir_path,
                                high_quality=high_quality,
                            )

                        if ok and pdf_path is not None and pdf_path.exists():
                            st.success("Conversion completed successfully.")
                            st.caption(message)
                            st.download_button(
                                label="Download PDF",
                                data=pdf_path.read_bytes(),
                                file_name=custom_pdf_name,
                                mime="application/pdf",
                                use_container_width=True,
                                key="download_pdf_hq",
                            )
                        else:
                            st.error(message)
                    except subprocess.TimeoutExpired:
                        st.error("Conversion timed out. Please try a smaller file.")
                    except Exception as exc:
                        st.error(f"Conversion failed: {exc}")

        st.markdown("</div>", unsafe_allow_html=True)

    with right3:
        st.markdown("<div class='soft-card'>", unsafe_allow_html=True)
        st.markdown("<div class='section-title'>Environment check</div>", unsafe_allow_html=True)
        if soffice_available:
            st.markdown(
                "<div class='status-neutral'>LibreOffice is installed and ready for high-quality DOCX → PDF export.</div>",
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                "<div class='status-warn'>LibreOffice is missing in the current environment. Add <b>libreoffice</b> to <b>packages.txt</b> for Streamlit deployment.</div>",
                unsafe_allow_html=True,
            )
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<div class='soft-card'>", unsafe_allow_html=True)
        st.markdown("<div class='section-title'>Best for</div>", unsafe_allow_html=True)
        st.markdown("<span class='blue-tag'>CV</span><span class='blue-tag'>Proposal</span><span class='blue-tag'>Report</span><span class='blue-tag'>Thesis</span>", unsafe_allow_html=True)
        st.markdown(
            "<div class='muted'>Choose this when you want a polished PDF with stronger image preservation from a DOCX file.</div>",
            unsafe_allow_html=True,
        )
        st.markdown("</div>", unsafe_allow_html=True)

# ==================================================
# NOTES
# ==================================================
with st.expander("Helpful notes"):
    st.markdown(
        """
- **PDF to Word (Editable)** is best when you want to modify text later.
- **PDF to Word (Preserve Appearance)** is best when you want the Word file to look as close to the original PDF as possible.
- **Word to PDF (High Quality)** keeps lossless compression and avoids image downsampling when LibreOffice is available.
- Put `requirements.txt` and `packages.txt` in the root of your GitHub repo.
- Files are processed in a temporary working folder and then cleaned up by the app logic.
        """
    )

st.markdown(
    """
    <div class="footer-note">
        Built with Streamlit for a cleaner and more professional document conversion experience.
    </div>
    """,
    unsafe_allow_html=True,
)
